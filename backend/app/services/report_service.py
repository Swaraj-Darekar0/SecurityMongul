import os
import json
import google.generativeai as genai
from flask import current_app
from docx import Document
from datetime import datetime

class ReportService:
    def __init__(self):
        self.data_dir = current_app.config['DATA_DIR']
        self.templates_dir = current_app.config['TEMPLATES_DIR']
        
        # Configure Gemini API
        genai.configure(api_key=os.getenv('GEMINI_API_KEY'))
        self.model = genai.GenerativeModel('gemini-1.5-pro')
    
    def generate_report(self, scan_id, report_type):
        """Generate report using Gemini with appropriate template"""
        try:
            # Load scan results
            scan_results = self._load_scan_results(scan_id)
            
            # Load appropriate template
            template = self._load_template(report_type)
            
            # Generate report content using Gemini
            report_content = self._generate_with_gemini(scan_results, template, report_type)
            
            # Format and save report
            report_path = self._format_and_save_report(report_content, scan_id, report_type)
            
            return report_path
            
        except Exception as e:
            raise Exception(f"Report generation failed: {str(e)}")
    
    def _load_template(self, report_type):
        """Load template based on report type"""
        template_mapping = {
            'regulatory_compliance': 'regulatory_compliance_template.md',
            'technical_operational': 'technical_operational_template.md',
            'business_focused': 'business_focused_template.md'
        }
        
        template_file = template_mapping.get(report_type)
        if not template_file:
            raise ValueError(f"Unknown report type: {report_type}")
        
        template_path = os.path.join(self.templates_dir, template_file)
        
        if not os.path.exists(template_path):
            raise FileNotFoundError(f"Template not found: {template_path}")
        
        with open(template_path, 'r', encoding='utf-8') as f:
            return f.read()
    
    def _generate_with_gemini(self, scan_results, template, report_type):
        """Use Gemini to generate report content"""
        prompt = f"""
        You are a professional security report generator. Use the provided template and security scan data to create a comprehensive {report_type.replace('_', ' ')} report.

        TEMPLATE:
        {template}

        SECURITY SCAN DATA:
        {json.dumps(scan_results, indent=2)}

        Instructions:
        1. Follow the template structure exactly
        2. Replace placeholders with actual data from the scan results
        3. Write in professional, clear language
        4. Include specific findings with file paths and line numbers where applicable
        5. Provide actionable recommendations
        6. Ensure compliance mappings are accurate
        
        Generate the complete report:
        """
        
        response = self.model.generate_content(prompt)
        return response.text
    
    def _load_scan_results(self, scan_id):
        """Load scan results from storage"""
        results_path = os.path.join(self.data_dir, 'scanned_results', f'{scan_id}.json')
        
        if not os.path.exists(results_path):
            raise FileNotFoundError(f"Scan results not found for ID: {scan_id}")
        
        with open(results_path, 'r') as f:
            return json.load(f)
    
    def _format_and_save_report(self, content, scan_id, report_type):
        """Format report content and save as document"""
        # Create reports directory if it doesn't exist
        reports_dir = os.path.join(self.data_dir, 'generated_reports')
        os.makedirs(reports_dir, exist_ok=True)
        
        # Generate filename
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"{report_type}_{scan_id}_{timestamp}.docx"
        report_path = os.path.join(reports_dir, filename)
        
        # Create Word document
        doc = Document()
        
        # Add title
        title = doc.add_heading(f'Security {report_type.replace("_", " ").title()} Report', 0)
        
        # Split content into paragraphs and add to document
        paragraphs = content.split('\n\n')
        for paragraph in paragraphs:
            if paragraph.strip():
                if paragraph.startswith('#'):
                    # Handle headers
                    level = paragraph.count('#')
                    text = paragraph.strip('#').strip()
                    doc.add_heading(text, level)
                else:
                    doc.add_paragraph(paragraph.strip())
        
        # Save document
        doc.save(report_path)
        return report_path
